"""Generate the Planner App project document as a .docx.

Mirrors the structure of the reference Stock-Analyst Google Doc:
  1. Project Outline   (visual overview / wireframes)
  2. Initial Prompt     (the original ask)
  3. Refined Prompt     (detailed, build-ready spec)
  4. Commands           (setup, run, git/github, sample usage)

Run with:  uv run --with python-docx python docs/generate_doc.py
The .docx imports straight into Google Docs (drag into Drive).
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

OUT = Path(__file__).resolve().parent / "Planner_App_Project_Doc.docx"

MONO = "Consolas"
CODE_SHADE = "F2F2F2"
GREEN = RGBColor(0x1B, 0x7F, 0x3B)


# --- low-level helpers -----------------------------------------------------

def shade(paragraph, fill: str) -> None:
    """Apply a solid background shade to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def code_block(doc: Document, text: str) -> None:
    """Render a monospace, shaded code/wireframe block (one paragraph)."""
    for i, line in enumerate(text.strip("\n").split("\n")):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        shade(p, CODE_SHADE)
        run = p.add_run(line if line else " ")
        run.font.name = MONO
        run.font.size = Pt(9)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def bullet(doc: Document, text: str, level: int = 0, mono_lead: str | None = None):
    style = "List Bullet" if level == 0 else f"List Bullet {level + 1}"
    p = doc.add_paragraph(style=style)
    if mono_lead:
        r = p.add_run(mono_lead)
        r.font.name = MONO
        r.font.color.rgb = GREEN
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def numbered(doc: Document, text: str):
    return doc.add_paragraph(text, style="List Number")


def h1(doc: Document, text: str, new_page: bool = True):
    if new_page:
        doc.add_page_break()
    doc.add_heading(text, level=1)


def h2(doc: Document, text: str):
    doc.add_heading(text, level=2)


def h3(doc: Document, text: str):
    doc.add_heading(text, level=3)


def para(doc: Document, text: str):
    return doc.add_paragraph(text)


# --- document content ------------------------------------------------------

def build() -> None:
    doc = Document()

    # Base font
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # ===== Title page =====
    title = doc.add_heading("🗓️  Daily Planner App", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Project Document  ·  Streamlit · uv · SQLite · Altair")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True
    note = doc.add_paragraph(
        "A four-tab planner where an event added in any tab appears across all of "
        "them, because every tab reads from one shared SQLite store."
    )
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # =========================================================
    # SECTION 1 — PROJECT OUTLINE
    # =========================================================
    h1(doc, "Project Outline", new_page=True)
    para(
        doc,
        "A personal daily-planning web app with four tabs: Daily, Weekly, Monthly, "
        "and Performance. All tabs share a single source of truth, so adding an "
        "event anywhere makes it appear everywhere. Visual overview of the "
        "application to be built:",
    )

    h2(doc, "Tab 1 — Daily")
    code_block(doc, r"""
+-----------------------------------------------+
| [Daily]  Weekly   Monthly   Performance       |
|              Thursday, June 4 2026            |
|  +-----------------------------------------+  |
|  |  Today's focus                          |  |
|  +-----------------------------------------+  |
|  +-----------------------------------------+  |
|  |  9:00 AM  . Team standup . Work  . 1hr  |  |
|  |  11:00 AM . Deep work    . Study . 2hr  |  |
|  |                                         |  |
|  |                + add event              |  |
|  +-----------------------------------------+  |
+-----------------------------------------------+
""")

    h2(doc, "Tab 2 — Weekly")
    code_block(doc, r"""
+-----------------------------------------------+
|  Daily  [Weekly]  Monthly   Performance       |
|  <        Thu, June 4 2026                >   |
|  Mo  Tu  We [Th] Fr  Sa  Su                   |
|   1   2   3   4   5   6   7                    |
|  ----------- Selected day detail -----------  |
|  |  9:00 AM . Team standup . Work . 1hr    |  |
|  |  2:00 PM . Client call  . Work . 1hr    |  |
|  |                + add event              |  |
+-----------------------------------------------+
""")

    h2(doc, "Tab 3 — Monthly")
    code_block(doc, r"""
+-----------------------------------------------+
|  Daily  Weekly  [Monthly]  Performance        |
|  <            June 2026                    >   |
|  Su  Mo  Tu  We  Th  Fr  Sa                   |
|       1   2   3  (4)  5   6     . = has events |
|   7   8   9  10  11  12  13                    |
|  14  15  16  17  18  19  20                    |
|  ...   click a day -> see what's planned       |
+-----------------------------------------------+
""")

    h2(doc, "Tab 4 — Performance")
    code_block(doc, r"""
+-----------------------------------------------+
|  Daily  Weekly  Monthly  [Performance]        |
|  [Today]  This week   This month              |
|            (  donut: time spent  )            |
|   Scheduled 6 hrs | Remaining 18 hrs | Total  |
|  --------- Category breakdown (bars) -------  |
|  Work  Study  Social  Hobby  Fun  Health      |
+-----------------------------------------------+
""")

    # =========================================================
    # SECTION 2 — INITIAL PROMPT
    # =========================================================
    h1(doc, "Initial Prompt")
    para(doc, "The original request that kicked off the build:")
    para(
        doc,
        "Build a daily planner web app with four tabs: Daily, Weekly, Monthly, and "
        "Performance. Any event added in one tab should automatically show up in the "
        "other tabs. Use Streamlit, and create a uv project.",
    )
    bullet(doc, "Tab 1 — Daily: add a “Today’s focus” note, plus a "
                "“+ add event” feature to add an event name and time. Added "
                "events are listed for the day.")
    bullet(doc, "Tab 2 — Weekly: show the current date at the top with left/right "
                "arrows to move to the previous/next week. Add an event to any day, "
                "including future days.")
    bullet(doc, "Tab 3 — Monthly: a full-month calendar with the ability to switch "
                "months and see what is planned for any given day / week / month.")
    bullet(doc, "Tab 4 — Performance: view performance for the day, week, and month. "
                "Classify events into categories (Work, Study, Social, Hobby, Fun, "
                "and others). Show a chart of time spent per category and the "
                "remaining time for each day / week / month.")
    bullet(doc, "Use the reference mockup image for the look and feel.")

    # =========================================================
    # SECTION 3 — REFINED PROMPT
    # =========================================================
    h1(doc, "Refined Prompt")

    h2(doc, "Core Requirements")
    para(doc, "Create a Streamlit app with exactly 4 tabs:")
    numbered(doc, "Daily")
    numbered(doc, "Weekly")
    numbered(doc, "Monthly")
    numbered(doc, "Performance")
    para(
        doc,
        "All four tabs read from and write to a single SQLite store, so an event "
        "added in any tab appears across all of them on the next rerun. Categories "
        "are a fixed, ordered set with stable colors: Work, Study, Social, Hobby, "
        "Fun, Health.",
    )

    h2(doc, "Data Model")
    para(doc, "A single SQLite database (data/planner.db) with two tables:")
    bullet(doc, "events (id, date, time, title, category, duration_hours)")
    bullet(doc, "focus (date, text) — one focus note per day")
    para(doc, "Every view is derived from the events table by filtering:")
    bullet(doc, "Daily — events where date == selected day")
    bullet(doc, "Weekly — events in the selected Monday–Sunday range")
    bullet(doc, "Monthly — events grouped by date (to draw the dots)")
    bullet(doc, "Performance — events grouped by category, durations summed")

    h2(doc, "Tab 1: Daily")
    bullet(doc, "Show the selected day’s long date as a centered header.")
    bullet(doc, "A “Today’s focus” text field, persisted per day.")
    bullet(doc, "List the day’s events as cards (time, title, category, duration) "
                "with a delete control.")
    bullet(doc, "A “+ add event” button opening a modal form.")

    h2(doc, "Tab 2: Weekly")
    bullet(doc, "Header shows the selected day with < and > arrows to move week by week.")
    bullet(doc, "Seven day buttons (Mon–Sun); a dot marks days that have events.")
    bullet(doc, "Clicking a day selects it and shows its events below.")
    bullet(doc, "“+ add event” can target any day, including future days.")

    h2(doc, "Tab 3: Monthly")
    bullet(doc, "Full-month grid (Sunday-first) with < and > month navigation.")
    bullet(doc, "Days with events are marked with a dot; today is highlighted.")
    bullet(doc, "Clicking a day shows what’s planned and allows adding events.")

    h2(doc, "Tab 4: Performance")
    bullet(doc, "Range toggle: Today / This week / This month.")
    bullet(doc, "Donut chart of time spent per category (Altair).")
    bullet(doc, "Metric cards: Scheduled, Remaining, Total hours.")
    bullet(doc, "Remaining = budget − scheduled, where budget = 24h/day "
                "(×7 for a week, ×days-in-month for a month).")
    bullet(doc, "Horizontal bar chart of the per-category breakdown.")
    bullet(doc, "Empty state when no events exist in the range.")

    h2(doc, "Required Project Structure")
    code_block(doc, r"""
Planner_App/
|-- app.py                     # entry point: page config, init, tab routing
|-- pyproject.toml             # uv project + dependencies
|-- uv.lock
|-- .python-version            # 3.12
|-- README.md
|-- .gitignore                 # ignores data/ and caches
|-- .streamlit/
|   `-- config.toml            # theme
|-- src/planner/
|   |-- __init__.py
|   |-- models.py              # Event dataclass, categories + colors
|   |-- storage.py             # SQLite layer (single source of truth)
|   |-- state.py               # session-state + date navigation
|   |-- components.py          # add-event dialog, event cards
|   `-- views/
|       |-- __init__.py
|       |-- daily.py
|       |-- weekly.py
|       |-- monthly.py
|       `-- performance.py
`-- data/planner.db            # local SQLite store (gitignored)
""")

    h2(doc, "File Responsibilities")
    h3(doc, "app.py")
    bullet(doc, "Main Streamlit entry point; set page config (title, icon, layout).")
    bullet(doc, "Initialize the database and session state.")
    bullet(doc, "Create the 4 tabs and call each view’s render() function.")
    h3(doc, "src/planner/storage.py")
    bullet(doc, "Own the SQLite connection (cached with @st.cache_resource).")
    bullet(doc, "CRUD for events; get/set the per-day focus note. Single source of truth.")
    h3(doc, "src/planner/models.py")
    bullet(doc, "Event dataclass and display helpers; CATEGORIES with stable colors.")
    h3(doc, "src/planner/state.py")
    bullet(doc, "Seed session_state once; date math (week bounds, shift week/month).")
    h3(doc, "src/planner/components.py")
    bullet(doc, "Reusable add-event modal dialog, event cards, and empty-state hints.")
    h3(doc, "src/planner/views/*.py")
    bullet(doc, "One module per tab; each exposes a render() that reads the shared store.")

    h2(doc, "uv Requirements")
    para(doc, "Use uv strictly for environment and package management. Setup commands:")
    code_block(doc, """
uv init --python 3.12
uv add streamlit altair pandas
uv run streamlit run app.py
""")
    para(doc, "Do not use pip, venv, conda, poetry, or requirements.txt. The project "
              "relies on pyproject.toml and uv.lock.")

    h2(doc, "Streamlit UI Requirements")
    bullet(doc, "Use st.set_page_config (centered layout, calendar icon).")
    bullet(doc, "Use st.tabs for the four tabs.")
    bullet(doc, "Use st.dialog for the add-event modal.")
    bullet(doc, "Namespace widget keys per view to avoid duplicate-key collisions.")
    bullet(doc, "After any write, call st.rerun() so all tabs reflect the change.")
    bullet(doc, "Theme via .streamlit/config.toml to match a clean, minimal look.")

    h2(doc, "README Requirements")
    para(doc, "Create a README.md containing:")
    numbered(doc, "Project overview")
    numbered(doc, "Tab-by-tab features")
    numbered(doc, "How to run (uv)")
    numbered(doc, "Thought process & design rationale (rerun model, single source of truth)")
    numbered(doc, "Project layout")
    numbered(doc, "Tech stack")

    h2(doc, "Notes & Simplifications")
    bullet(doc, "Single local user; no authentication.")
    bullet(doc, "No external API keys or environment variables are required.")
    bullet(doc, "“Remaining” uses a flat 24h/day budget rather than waking hours.")
    bullet(doc, "Editing an event = delete + re-add (no in-place edit form yet).")

    # =========================================================
    # SECTION 4 — COMMANDS
    # =========================================================
    h1(doc, "Commands")

    h2(doc, "Project setup (uv)")
    code_block(doc, """
# Install uv (macOS / Linux)
curl -LsSf https://astral.sh/uv/install.sh | sh

# Initialize the project and add dependencies
uv init --python 3.12
uv add streamlit altair pandas
""")

    h2(doc, "Run the app")
    code_block(doc, """
uv run streamlit run app.py
# opens http://localhost:8501
""")

    h2(doc, "Version control / GitHub")
    code_block(doc, """
# Install GitHub CLI and authenticate
gh auth login            # GitHub.com -> HTTPS -> web browser

# Add this app as a subfolder of an existing repo
gh repo clone <owner>/AgenticAI
cp -R Planner_App AgenticAI/Planner_App
cd AgenticAI
git add Planner_App && git commit -m "Add Planner_App" && git push
""")

    h2(doc, "Sample usage")
    bullet(doc, "Daily tab: type a focus, click “+ add event”, fill name/time/"
                "category/duration, Add.")
    bullet(doc, "Weekly tab: use < > to move weeks, click a day, add a future event.")
    bullet(doc, "Monthly tab: switch months with < >, click a dotted day to see its plan.")
    bullet(doc, "Performance tab: toggle Today / This week / This month to see the donut "
                "and remaining hours update.")

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
